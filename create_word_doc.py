import os
import sys

# Ensure python-docx is installed
try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    print("Installing python-docx library...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)
        
    # Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2C, 0x2D, 0x30) # Charcoal text
    
    # 1. Project Title
    title = doc.add_paragraph()
    title_run = title.add_run("CommunityHero AI")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1C, 0x1D, 0x1F)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run("Empowering Communities Through AI-Driven Civic Action")
    tagline_run.italic = True
    tagline_run.font.size = Pt(12)
    tagline_run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC) # Deep blue
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # Helper function for headings
    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1C, 0x1D, 0x1F)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    # 2. Problem Statement
    add_heading("1. Problem Statement")
    p1 = doc.add_paragraph("Municipal infrastructure management is slow, non-transparent, and disconnected from the everyday experiences of citizens. When local hazards—such as dangerous potholes, overflowing garbage containers, sewage leaks, or broken streetlights—go unreported, they lead to public safety risks, environmental degradation, and neighborhood decline. Traditional reporting portals are cumbersome, lack user engagement, and provide no real-time progress indicators, resulting in low resident participation.")
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(10)

    # 3. Solution
    add_heading("2. The Solution")
    p2 = doc.add_paragraph("CommunityHero AI turns city operations into an intelligent civic operating system. It is a full-stack platform that empowers citizens to act as first responders. Citizens upload a photo of a hazard; an AI computer vision engine (YOLOv8) automatically identifies the category, confidence rate, and severity level. Local coordinates are verified by neighbors on interactive maps to prevent spam. Public works authorities assign dispatches and verify repairs using before/after visual similarity checks (SSIM) to close tasks and award citizen points.")
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(10)

    # 4. Key Features
    add_heading("3. Key Features")
    
    # Feature 1
    f1 = doc.add_paragraph()
    f1_title = f1.add_run("• AI Computer Vision: ")
    f1_title.bold = True
    f1_desc = f1.add_run("A FastAPI Python microservice running a pretrained YOLOv8 deep learning model. It performs real-time visual scans on uploaded photographs to detect potholes, trash overflow, or damaged utilities with high confidence scores, setting priority levels automatically.")
    
    # Feature 2
    f2 = doc.add_paragraph()
    f2_title = f2.add_run("• Live Viewport Rendering: ")
    f2_title.bold = True
    f2_desc = f2.add_run("An interactive, high-frequency coordinate mapping interface built with Leaflet and OpenStreetMap Nominatim APIs. As users search or pan anywhere globally, the viewport boundaries dynamically generate local issues and coordinate pin metrics.")
    
    # Feature 3
    f3 = doc.add_paragraph()
    f3_title = f3.add_run("• Hero Point Gamification: ")
    f3_title.bold = True
    f3_desc = f3.add_run("A reputation scoring engine that credits points (+15 for reporting, +5 for verifying, +50 for resolution) and unlocks achievement profile badges. This turns civic cleanup into an engaging community leaderboard experience.")

    # 5. Tech Stack
    add_heading("4. Tech Stack")
    t1 = doc.add_paragraph()
    t1.add_run("• Frontend: ").bold = True
    t1.add_run("Next.js 14 App Router, Tailwind CSS, Zustand, Framer Motion, Leaflet Maps")
    
    t2 = doc.add_paragraph()
    t2.add_run("• Backend Gateway: ").bold = True
    t2.add_run("Node.js, Express.js, MongoDB (Mongoose), Socket.IO WebSockets")
    
    t3 = doc.add_paragraph()
    t3.add_run("• AI Microservice: ").bold = True
    t3.add_run("FastAPI Python 3.10, OpenCV (Grayscale SSIM), YOLOv8 (Ultralytics)")

    # 6. How to Run / Use
    add_heading("5. Setup & Execution Guide")
    
    r1 = doc.add_paragraph("To run the entire system locally using Docker Compose:")
    r1.paragraph_format.space_after = Pt(4)
    
    c1 = doc.add_paragraph()
    c1_run = c1.add_run("   docker-compose up --build")
    c1_run.font.name = 'Courier New'
    c1_run.font.size = Pt(10)
    c1.paragraph_format.space_after = Pt(6)
    
    r2 = doc.add_paragraph("Access the services on your local ports:")
    r2.paragraph_format.space_after = Pt(4)
    
    c2 = doc.add_paragraph()
    c2_run = c2.add_run("   - Next.js Client: http://localhost:3000\n   - Express API Server: http://localhost:5000\n   - FastAPI AI Service: http://localhost:8000")
    c2_run.font.name = 'Courier New'
    c2_run.font.size = Pt(9.5)
    
    # Save document
    filename = "CommunityHero_AI_Project_Description.docx"
    doc.save(filename)
    print(f"Document created successfully: {filename}")

if __name__ == '__main__':
    create_document()
